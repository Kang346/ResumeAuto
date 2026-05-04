#!/usr/bin/env python3
"""Extract plain text from a .docx resume to stdout.

Usage:
    python tools/extract_resume_text.py path/to/resume.docx
"""
import sys
from pathlib import Path


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: extract_resume_text.py <path-to-docx>", file=sys.stderr)
        return 2
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"file not found: {path}", file=sys.stderr)
        return 1
    if path.suffix.lower() != ".docx":
        print(f"only .docx is supported (got {path.suffix})", file=sys.stderr)
        return 1
    from docx import Document
    doc = Document(str(path))
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            print(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                print(" | ".join(cells))
    return 0


if __name__ == "__main__":
    sys.exit(main())
